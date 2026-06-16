import os
import pdfplumber
import docx

def extract_text_from_pdf(file_path: str) -> str:
    """
    Extracts complete text from a PDF file using pdfplumber.
    Also extracts any embedded hyperlinks and appends them to the end of the text.
    Raises ValueError if text extraction is entirely empty or fails.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found at: {file_path}")
        
    text_content = []
    hyperlinks = []
    try:
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text(x_tolerance=2)
                if page_text:
                    text_content.append(page_text)
                
                # Extract embedded hyperlinks / annotations
                try:
                    if hasattr(page, "hyperlinks") and page.hyperlinks:
                        for hl in page.hyperlinks:
                            if isinstance(hl, dict) and hl.get("uri"):
                                hyperlinks.append(hl["uri"].strip())
                except Exception:
                    pass # Ignore errors in annot extraction
    except Exception as e:
        raise ValueError(f"Failed to parse PDF document structure. It might be corrupt: {str(e)}")
        
    extracted_text = "\n".join(text_content).strip()
    if not extracted_text:
        raise ValueError("The PDF file seems to be empty or contains only non-extractable scanned images.")
        
    if hyperlinks:
        # De-duplicate links while preserving order
        seen_links = set()
        unique_links = []
        for link in hyperlinks:
            if link not in seen_links:
                seen_links.add(link)
                unique_links.append(link)
        
        extracted_text += "\n\n--- Extracted Hyperlinks ---\n" + "\n".join(unique_links)
        
    return extracted_text

def extract_text_from_docx(file_path: str) -> str:
    """
    Extracts complete text from a DOCX file, including both paragraphs
    and table contents, ensuring structured elements are captured.
    Raises ValueError if extraction fails or is empty.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found at: {file_path}")
        
    text_content = []
    try:
        doc = docx.Document(file_path)
        
        # 1. Extract paragraph content
        for para in doc.paragraphs:
            val = para.text.strip()
            if val:
                text_content.append(val)
                
        # 2. Extract table content (essential as resumes often use tables for layouts)
        for table in doc.tables:
            for row in table.rows:
                # Merge duplicate cell values due to merged cells
                row_cells_text = []
                for cell in row.cells:
                    c_text = cell.text.strip()
                    if c_text and (not row_cells_text or row_cells_text[-1] != c_text):
                        row_cells_text.append(c_text)
                if row_cells_text:
                    text_content.append(" | ".join(row_cells_text))
                    
    except Exception as e:
        raise ValueError(f"Failed to parse Word document structure. It might be corrupt: {str(e)}")
        
    extracted_text = "\n".join(text_content).strip()
    if not extracted_text:
        raise ValueError("The DOCX file seems to be empty.")
        
    return extracted_text
