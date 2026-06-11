import os
import docx

def create_docx_resume():
    doc = docx.Document()
    
    # Header Name
    h = doc.add_heading("Kavya Sri", level=0)
    h.alignment = 1 # Center alignment
    
    # Contact Row
    p_contact = doc.add_paragraph()
    p_contact.alignment = 1
    p_contact.add_run("Email: ").bold = True
    p_contact.add_run("kavya@gmail.com | ")
    p_contact.add_run("Phone: ").bold = True
    p_contact.add_run("+91-9876543210 | ")
    p_contact.add_run("LinkedIn: ").bold = True
    p_contact.add_run("https://www.linkedin.com/in/kavyasri121 | ")
    p_contact.add_run("GitHub: ").bold = True
    p_contact.add_run("https://github.com/kavyasri121")
    
    # Summary
    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph("Results-driven Software Engineer with 3+ years of experience designing and implementing highly scalable FastAPI backend web applications and responsive React user interfaces.")

    # Skills
    doc.add_heading("Technical Skills", level=1)
    doc.add_paragraph("Languages: Python, JavaScript, SQL, HTML5, CSS3")
    doc.add_paragraph("Frameworks & Tools: FastAPI, React, Node.js, Git, Docker, AWS")

    # Experience
    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph("Software Engineer | CloudTech Solutions (2024 - Present)")
    doc.add_paragraph("• Led backend migration to FastAPI, reducing response latencies by 35%.")
    doc.add_paragraph("• Structured multi-tenant SQL database architectures utilizing PostgreSQL.")
    
    doc.add_paragraph("Junior Web Developer | Innovate Softwares (2022 - 2024)")
    doc.add_paragraph("• Implemented responsive UI mockups in HTML/CSS/React matching UI/UX design specifications.")
    doc.add_paragraph("• Maintained RESTful APIs and performed regular unit tests using PyTest.")

    # Education
    doc.add_heading("Education", level=1)
    doc.add_paragraph("Bachelor of Technology in Computer Science & Engineering")
    doc.add_paragraph("ABC Institute of Technology | CGPA: 9.2 (2018 - 2022)")

    # Projects
    doc.add_heading("Projects", level=1)
    doc.add_paragraph("AI Resume Analyzer: Programmed a secure text extraction microservice parsing PDF and Word files with FastAPI.")
    doc.add_paragraph("Task Planner Dashboard: Developed a drag-and-drop task scheduler using React and Node.js.")

    # Certifications
    doc.add_heading("Certifications", level=1)
    doc.add_paragraph("• AWS Certified Cloud Practitioner")
    doc.add_paragraph("• Certified JavaScript Developer - W3C")

    # Languages
    doc.add_heading("Languages", level=1)
    doc.add_paragraph("English (Professional), Telugu (Native), Hindi (Conversational)")

    file_path = os.path.join(os.path.dirname(__file__), "test_resume.docx")
    doc.save(file_path)
    print(f"Successfully created mock resume at: {file_path}")

if __name__ == "__main__":
    create_docx_resume()
